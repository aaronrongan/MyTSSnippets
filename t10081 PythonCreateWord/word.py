# print("Hello world")

# read and save word doc
from docx import Document

# open an document
doc = Document('C:\\Users\\Aaron\\Documents\\MyZoteroFiles_PNF\\PNF.docx')

doc.add_heading("PNF21061501 设计文档",level=0)


doc.add_heading('目的',level=1)
doc.add_paragraph(' ', style='List Bullet')   #插入段落
doc.add_paragraph(' ', style='Normal')   #插入段落

# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

doc.add_heading('参考文献',level=1)
doc.add_paragraph(' ', style='List Number w Bracket')
doc.add_paragraph(' ', style='Normal')   #插入段落

doc.add_heading('日志',level=1)
table = doc.add_table(rows=3, cols=3,style='Table Grid') #插入表格
# table = doc.add_table(rows=3, cols=3,style='Plain Table 2') #插入表格
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '日期'
hdr_cells[1].text = '内容'
hdr_cells[2].text = '总结'

doc.add_heading('问题',level=1)
p = doc.add_paragraph('',style='List Bullet')   #插入段落
doc.add_paragraph(' ', style='Normal')   #插入段落

doc.add_heading('笔记',level=1)
p = doc.add_paragraph('',style='List Bullet')   #插入段落
doc.add_paragraph(' ', style='Normal')   #插入段落

# 保存一个文件
doc.save('C:\\Users\\Aaron\\Documents\\MyZoteroFiles_PNF\\a.docx')