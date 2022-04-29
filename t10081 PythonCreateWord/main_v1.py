#-*- coding: utf-8 -*-

import wx
import win32api
import sys, os

# read and save word doc
from docx import Document

import datetime
import webbrowser
# from Tkinter import *

from wx.core import MessageBox

now_time = datetime.datetime.now()


APP_TITLE = u'文档自动生成'
APP_ICON = 'D:/Program Files/Microsoft VS Code/resources/app/resources/win32/python.ico' # 请更换成你的icon

Path_DHF='C:/Users/Aaron/Documents/MyZoteroFiles_DHF'
Path_EDU='C:/Users/Aaron/Documents/MyZoteroFiles_EDU'
Path_IRF='C:/Users/Aaron/Documents/MyZoteroFiles_IRF'
Path_PNF='C:/Users/Aaron/Documents/MyZoteroFiles_PNF'
Path_RNF='C:/Users/Aaron/Documents/MyZoteroFiles_RNF'
Path_SHF='C:/Users/Aaron/Documents/MyZoteroFiles_SHF'

gFileName=''
gTodayText=datetime.datetime.now().strftime('%y%m%d')



class mainFrame(wx.Frame):
    '''程序主窗口类，继承自wx.Frame'''
    
    def __init__(self, parent):
        '''构造函数'''
        
        wx.Frame.__init__(self, parent, -1, APP_TITLE)
        self.SetBackgroundColour(wx.Colour(250,250,250))
        self.SetSize((520, 190))
        self.Center()
        
        if hasattr(sys, "frozen") and getattr(sys, "frozen") == "windows_exe":
            exeName = win32api.GetModuleFileName(win32api.GetModuleHandle(None))
            icon = wx.Icon(exeName, wx.BITMAP_TYPE_ICO)
        else :
            icon = wx.Icon(APP_ICON, wx.BITMAP_TYPE_ICO)
        self.SetIcon(icon)
        
       
        # FileFormat=[('Docx',0),('Txt',1),('Xlsx',2)]

        # pnl = wx.Panel(self) 
        wx.StaticText(self, -1, u'文件类型：', pos=(10,15), size=(80, -1), style=wx.ALIGN_LEFT)
        # root_filetype=Tk()
        pnl = wx.Panel(self)
        self.rb1 = wx.RadioButton(self, label = 'DHF',pos = (10,30)) 
        self.rb2 = wx.RadioButton(self, label = 'EDU',pos = (10,45)) 
        self.rb3 = wx.RadioButton(self, label = 'IRF',pos = (10,60)) 
        self.rb4 = wx.RadioButton(self, label = 'PNF',pos = (10,75)) 
        self.rb5 = wx.RadioButton(self, label = 'RNF',pos = (10,90)) 
        self.rb6 = wx.RadioButton(self, label = 'SHF',pos = (10,105)) 

        # FileTypeList=[('DHF',0),('EDU',1),('IRF',2),('PNF',3),('RNF',4),('SHF',5)]
        # FileTypeList = ['DHF', 'EDU','IRF','PNF','RNF','SHF']
        # self.rbox = wx.RadioBox(pnl, label = 'RadioBox', pos = (10,10), choices = FileTypeList,
        #     majorDimension = 1, style = wx.RA_SPECIFY_ROWS)
        # self.rbox.Bind(wx.EVT_RADIOBOX,self.onRadioBox)

        # for  ft,num in FileType:
        #     # Radiobutton(root, text=lan, value=num, command=callRB, variable=v,indicatoron=0,width=10,anchor='c').pack()
        #     # Radiobutton(root, text=lan, value=num, command=callRB, variable=v,indicatoron=0,width=10,anchor='c').pack()
        #     Radiobutton(root, text=ft, value=num, command=callRB, variable=v).pack(anchor=W)

        # self.Bind(wx.EVT_CHECKBOX,self.on) 
        self.Centre() 
        self.Show(True)

        # self.cb2 = wx.CheckBox(pnl, label = 'Value B',pos = (10,40)) 
        # self.cb3 = wx.CheckBox(pnl, label = 'Value C',pos = (10,70)) 

        wx.StaticText(self, -1, u'文件编号：', pos=(80, 15), size=(100, -1), style=wx.ALIGN_RIGHT)
        self.tc1 = wx.TextCtrl(self, -1, '', pos=(180, 10), size=(80, -1), name='TC01', style=wx.ALIGN_RIGHT)
        
        wx.StaticText(self, -1, u'文件名：', pos=(260, 15), size=(100, -1), style=wx.ALIGN_RIGHT)
        # self.tc2 = wx.TextCtrl(self, -1, '', pos=(145, 110), size=(150, -1), name='TC02', style=wx.TE_PASSWORD|wx.ALIGN_RIGHT)
        self.tc2 = wx.TextCtrl(self, -1, '', pos=(360, 10), size=(120, -1), name='TC02', style=wx.ALIGN_RIGHT)

        # self.tip = wx.StaticText(self, -1, u'', pos=(185, 150), size=(150, -1), style=wx.ST_NO_AUTORESIZE)
        wx.StaticText(self, -1, u'文件全名：', pos=(80, 50), size=(100, -1), style=wx.ALIGN_RIGHT)
        self.FullFileName = wx.StaticText(self, -1, u'hello', pos=(180, 50), size=(300, -1), style=wx.ST_NO_AUTORESIZE)

        # wx.StaticText(self, -1, u'文件格式：', pos=(80,80), size=(100, -1), style=wx.ALIGN_RIGHT)
        # self.rb_word = wx.RadioButton(self, label = 'Word',pos = (180,80), size=(50,-1),style=wx.ALIGN_RIGHT) 
        # self.rb_txt = wx.RadioButton(self, label = 'Txt',pos = (270,80), size=(50,-1),style=wx.ALIGN_RIGHT) 
        
        FileFormatList = ['Word', 'Txt']
        self.rbox = wx.RadioBox(pnl, label = 'RadioBox', pos = (80,100), choices = FileFormatList,
            majorDimension = 1, style = wx.RA_SPECIFY_ROWS)
        self.rbox.Bind(wx.EVT_RADIOBOX,self.onRadioBox)
        
        
        # btn_mea = wx.Button(self, -1, u'鼠标左键事件', pos=(80, 150), size=(100, 25))
        # btn_meb = wx.Button(self, -1, u'鼠标所有事件', pos=(350, 150), size=(100, 25))
        btn_createfile = wx.Button(self, -1, u'生成文件', pos=(150, 110), size=(100, 30))
        btn_close = wx.Button(self, -1, u'关闭窗口', pos=(300, 110), size=(100, 30))
        self.Bind(wx.EVT_BUTTON, self.OnClose, btn_close)
        
        # 控件事件
        self.tc1.Bind(wx.EVT_TEXT, self.EvtText)
        self.tc2.Bind(wx.EVT_TEXT, self.EvtText)

        self.rb1.Bind(wx.EVT_TEXT, self.EvtText)
        self.rb2.Bind(wx.EVT_TEXT, self.EvtText)
        self.rb3.Bind(wx.EVT_TEXT, self.EvtText)
        self.rb4.Bind(wx.EVT_TEXT, self.EvtText)
        self.rb5.Bind(wx.EVT_TEXT, self.EvtText)
        self.rb6.Bind(wx.EVT_TEXT, self.EvtText)

        # self.Bind(wx.EVT_BUTTON, self.OnClose, btn_close)
        self.Bind(wx.EVT_BUTTON,self.OnWriteDoc,btn_createfile)
        
        # 鼠标事件 
        # btn_mea.Bind(wx.EVT_LEFT_DOWN, self.OnLeftDown)
        # btn_mea.Bind(wx.EVT_LEFT_UP, self.OnLeftUp)
        # btn_mea.Bind(wx.EVT_MOUSEWHEEL, self.OnMouseWheel)
        # btn_meb.Bind(wx.EVT_MOUSE_EVENTS, self.OnMouse)
        
        # 键盘事件
        self.Bind(wx.EVT_KEY_DOWN, self.OnKeyDown)
        
        # 系统事件
        self.Bind(wx.EVT_CLOSE, self.OnClose)
        self.Bind(wx.EVT_SIZE, self.On_size)
        #self.Bind(wx.EVT_PAINT, self.On_paint)
        #self.Bind(wx.EVT_ERASE_BACKGROUND, lambda event: None)

        self.tc1.Value=gTodayText
    
    # 获取当前日期，以210615格式
    def GetCurrentYMD():
        pass

    def onRadioBox(self,e):
        pass

    def EvtText(self, evt):
        '''输入框事件函数'''
        
        obj = evt.GetEventObject()
        objName = obj.GetName()

        text = evt.GetString()
        TypeText='DHF'

        if self.rb1.Value==True:
            TypeText='DHF'
            # MessageBox('ok')
        elif self.rb2.Value==True:
            TypeText='EDU'
        elif self.rb3.Value==True:
            TypeText='IRF'
        elif self.rb4.Value==True:
            TypeText='PNF'
        elif self.rb5.Value==True:
            TypeText='RNF'
        elif self.rb6.Value==True:
            TypeText='SHF'

        NumberText=self.tc1.GetValue()
        NameText=self.tc2.GetValue()

        FullFileNameText=TypeText + NumberText + " " + NameText + ".docx"

        if objName == 'TC01':
            self.FullFileName.SetLabelText(FullFileNameText)


        elif objName == 'TC02':
            self.FullFileName.SetLabelText(FullFileNameText)
            
    
    def On_size(self, evt):
        '''改变窗口大小事件函数'''
        
        self.Refresh()
        evt.Skip() # 体会作用
    
    def OnClose(self, evt):
        '''关闭窗口事件函数'''
        
        # dlg = wx.MessageDialog(None, u'确定要关闭本窗口？', u'操作提示', wx.YES_NO | wx.ICON_QUESTION)
        # if(dlg.ShowModal() == wx.ID_YES):
        self.Destroy()
    
    def OnLeftDown(self, evt):
        '''左键按下事件函数'''
        
        # self.tip.SetLabel(u'左键按下')
        # gFileName=self.tc1.GetLabelText & self.tc2.GetLabelText &  ".docx"
        # self.FullFileName.SetLabel(gFileName)
        pass
    
    def OnLeftUp(self, evt):
        '''左键弹起事件函数'''
        
        # self.tip.SetLabel(u'左键弹起')
        # gFileName=self.tc1.GetLabelText & self.tc2.GetLabelText &  ".docx"
        # self.FullFileName.SetLabel(gFileName)
        pass
    
    def OnMouseWheel(self, evt):
        '''鼠标滚轮事件函数'''
        
        vector = evt.GetWheelRotation()
        # self.tip.SetLabel(str(vector))
        # self.FullFileName.SetLabel(str(vector))
    
    def OnMouse(self, evt):
        '''鼠标事件函数'''
        pass
        # self.tip.SetLabel(str(evt.EventType))
        

        # self.FullFileName.SetLabel(gFileName)
    
    def OnKeyDown(self, evt):
        '''键盘事件函数'''
        
        # key = evt.GetKeyCode() 
        # self.tip.SetLabel(str(key))
        self.FullFileName.SetLabel(u'KeyDown')
    
    # 保存文件():
    # def SaveFile():
    #     WriteDocTitle()
        
    # 打开文件，写入标题
    def OnWriteDoc(self, evt):
        

        if self.rb1.Value==True:
            FileType='DHF'
            # MessageBox('ok')
        elif self.rb2.Value==True:
            FileType='EDU'
        elif self.rb3.Value==True:
            FileType='IRF'
        elif self.rb4.Value==True:
            FileType='PNF'
        elif self.rb5.Value==True:
            FileType='RNF'
        elif self.rb6.Value==True:
            FileType='SHF'
        
        # open an document
        TemplateFileFullPathName='C:/Users/Aaron/Documents/MyZoteroFiles_' + FileType + '/'+ FileType+'.docx'

        # MessageBox(TemplateFileFullPathName)
        # FileNumber=
        ActualFileName=self.FullFileName.GetLabel()
        # MessageBox(ActualFileName)
        # open an document 
        doc = Document(TemplateFileFullPathName)

        # 应该把.docx后缀删除
        
        doc.add_heading(ActualFileName,level=0)

        doc.add_heading('目的',level=1)
        doc.add_paragraph(' ', style='List Bullet')   #插入段落
        doc.add_paragraph(' ', style='Normal')   #插入段落

        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True

        doc.add_heading('参考文献',level=1)
        doc.add_paragraph(' ', style='List Number w Bracket')
        doc.add_paragraph(' ', style='Normal')   #插入段落

        doc.add_heading('日志',level=1)
        table = doc.add_table(rows=3, cols=3,style='Table Grid') #插入表格
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '日期'
        hdr_cells[1].text = '内容'
        hdr_cells[2].text = '总结'

        doc.add_heading('问题',level=1)
        p = doc.add_paragraph('',style='List Bullet')   #插入段落
        doc.add_paragraph(' ', style='Normal')   #插入段落

        doc.add_heading('笔记',level=1)
        p = doc.add_paragraph('',style='List Bullet')   #插入段落
        doc.add_paragraph(' ', style='Normal')   #插入段落
        # 保存一个文件
        
        doc.save('C:/Users/Aaron/Documents/MyZoteroFiles_' + FileType + '/'+ ActualFileName)
        webbrowser.open('C:/Users/Aaron/Documents/MyZoteroFiles_' + FileType + '/'+ ActualFileName)

        
class mainApp(wx.App):
    def OnInit(self):
        self.SetAppName(APP_TITLE)
        self.Frame = mainFrame(None)
        self.Frame.Show()
        
        return True

if __name__ == "__main__":
    app = mainApp()
    app.MainLoop()
